"""Report exporters — PDF and DOCX generation from Markdown content."""

from __future__ import annotations

import io
import logging
import re

logger = logging.getLogger(__name__)


def export_to_pdf(title: str, html_content: str) -> bytes:
    """Convert styled HTML to PDF bytes using WeasyPrint.

    Args:
        title: Report title for PDF metadata.
        html_content: Full HTML string (from _wrap_html).

    Returns:
        PDF file as bytes.
    """
    try:
        from weasyprint import HTML

        pdf = HTML(string=html_content).write_pdf()
        return pdf
    except ImportError:
        logger.error("weasyprint not installed — cannot generate PDF")
        raise
    except Exception as e:
        logger.error("PDF generation failed: %s", e)
        raise


def export_to_docx(title: str, markdown_content: str) -> bytes:
    """Convert Markdown content to DOCX bytes using python-docx.

    Parses Markdown headings, tables, code blocks, and paragraphs
    to build a structured Word document.

    Args:
        title: Document title.
        markdown_content: Report in Markdown format.

    Returns:
        DOCX file as bytes.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx not installed — cannot generate DOCX")
        raise

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # Title
    doc.add_heading(title, level=0)

    lines = markdown_content.split("\n")
    i = 0
    in_table = False
    table_rows = []
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.startswith("```"):
            if in_code:
                # End code block
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run("\n".join(code_lines))
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if "|" in line and line.strip().startswith("|"):
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            # Skip separator rows
            if all(set(c) <= set("-: ") for c in cells):
                i += 1
                continue
            table_rows.append(cells)
            # Check if next line is still a table row
            if i + 1 < len(lines) and "|" in lines[i + 1] and lines[i + 1].strip().startswith("|"):
                i += 1
                continue
            else:
                # End of table — flush
                if table_rows:
                    cols = max(len(r) for r in table_rows)
                    table = doc.add_table(rows=len(table_rows), cols=cols)
                    table.style = "Table Grid"
                    for ri, row in enumerate(table_rows):
                        for ci, cell in enumerate(row):
                            if ci < cols:
                                table.rows[ri].cells[ci].text = cell
                table_rows = []
                i += 1
                continue

        # Flush any pending table
        if table_rows:
            cols = max(len(r) for r in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=cols)
            table.style = "Table Grid"
            for ri, row in enumerate(table_rows):
                for ci, cell in enumerate(row):
                    if ci < cols:
                        table.rows[ri].cells[ci].text = cell
            table_rows = []

        # Headings
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:])
            p.paragraph_format.left_indent = Inches(0.3)
        elif line.startswith("---"):
            doc.add_paragraph("─" * 50)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip() == "":
            pass  # Skip blank lines
        else:
            # Inline formatting: bold and code
            text = line
            p = doc.add_paragraph()
            # Split by bold and code markers
            parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith("`") and part.endswith("`"):
                    run = p.add_run(part[1:-1])
                    run.font.name = "Consolas"
                    run.font.size = Pt(9)
                else:
                    p.add_run(part)

        i += 1

    # Flush any remaining table
    if table_rows:
        cols = max(len(r) for r in table_rows)
        table = doc.add_table(rows=len(table_rows), cols=cols)
        table.style = "Table Grid"
        for ri, row in enumerate(table_rows):
            for ci, cell in enumerate(row):
                if ci < cols:
                    table.rows[ri].cells[ci].text = cell

    # Write to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
